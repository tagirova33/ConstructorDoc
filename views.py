from django.shortcuts import render, get_object_or_404
from .models import Person
from django.views.generic import View
from django.http import HttpResponse
from constructor.templates.co.forms import PersonForm

from django.http import FileResponse

import docx
from docxtpl import DocxTemplate

def index(request):
    return HttpResponse("Hello World!")

class ConsCreate(View):

    def get(self, request):
        form = PersonForm()
        return render(request, 'co/person_create.html', context={'form': form})

    def post(self, request):
        bound_form = PersonForm(request.POST)
        if bound_form.is_valid():
            var5 =  DocxTemplate("variable5.docx")
            var3 =  DocxTemplate("variable3.docx")
            var7 =  DocxTemplate("variable7.docx")
            basedock = DocxTemplate("basedock.docx")
            #final_dock = DocxTemplate("/Users/liliatagirova/Desktop/Доки/final.docx")

            myString3 = ""
            myString5 = ""
            myString7 = ""
            first_name = request.POST.get("first_name")
            last_name = request.POST.get("last_name")
            fam_name = request.POST.get("fam_name")

            option = request.POST.getlist('option')
            for i in range(len(option)):
                if option[i] == '3':
                    text3 = []
                    for paragraph in var3.paragraphs:
                        text3.append(paragraph.text)
                    myString3 = '\n\t'.join(text3)
                if option[i] == '5':
                    text5 = []
                    for paragraph in var5.paragraphs:
                        text5.append(paragraph.text)
                    myString5 = '\n\t'.join(text5)
                if option[i] == '7':
                    text7 = []
                    for paragraph in var7.paragraphs:
                        text7.append(paragraph.text)
                    myString7 = '\n\t'.join(text7)

            context = {'first_name' : first_name,'last_name': last_name, 'fam_name': fam_name,
            'third' : myString3, 'fifth' : myString5, 'seventh' : myString7
            }

            basedock.render(context)
            basedock.save("final.docx")

        return FileResponse(open("final.docx", 'rb'))
