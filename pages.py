# -*- coding: utf-8 -*-
import docx
from docxtpl import DocxTemplate
from constructor.templates.co.forms import PersonForm

basedock = DocxTemplate("/Users/liliatagirova/Desktop/Доки/basedock.docx")
var3 =  DocxTemplate("/Users/liliatagirova/Desktop/Доки/variable3.docx")
var5 =  DocxTemplate("/Users/liliatagirova/Desktop/Доки/variable5.docx")
var7 =  DocxTemplate("/Users/liliatagirova/Desktop/Доки/variable7.docx")
final_dock = DocxTemplate("final.docx")
name = PersonForm.first_name
print(name)
context = { 'name' : name, 'last_name': " 11", 'fam_name': "22"}

'''
doc1 = docx.Document('example.docx')
doc2 = docx.Document('restyled.docx')

# получаем из первого документа стили всех абзацев
styles = []
for paragraph in doc1.paragraphs:
    styles.append(paragraph.style)

# применяем стили ко всем абзацам второго документа
for i in range(len(doc2.paragraphs)):
    doc2.paragraphs[i].style = styles[i]

doc2.save('restored.docx')
'''
text = []
for paragraph in var5.paragraphs:
    text.append(paragraph.text)
#print('\n'.join(text))
myString = '/n/t'.join(text)
#context = { 'third' : myString}
context = { 'fifth' : myString}
#context = { 'seventh' : myString}
final_dock.render(context)
final_dock.save("final.docx")
print("true")
